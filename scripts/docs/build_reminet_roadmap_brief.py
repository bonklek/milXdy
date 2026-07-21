from __future__ import annotations

import re
import sys
from pathlib import Path
import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "ideas" / "meeting-drafts" / "milxdy-roadmap-brief-reminet-2026-07-20.md"
OUTPUT = Path(
    os.environ.get(
        "MILXDY_ROADMAP_DOCX_OUTPUT",
        ROOT / "ideas" / "meeting-drafts" / "milxdy-roadmap-brief-reminet-working-draft.docx",
    )
)

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "59636E"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF8E8"
BORDER = "CBD2D9"
WHITE = "FFFFFF"
BLACK = "111111"

CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, name="Calibri", size=None, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_borders(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_borders(cell)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([r_pr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(r"(\[[^\]]+\]\([^)]+\)|\*\*[^*]+\*\*|`[^`]+`)")


def add_inline(paragraph, text: str, default_color=BLACK, default_size=11):
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=default_size, color=default_color)
        token = match.group(0)
        if token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=default_size, color=default_color, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=default_size - 0.5, color=DARK_BLUE)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=default_size, color=default_color)


def clean_markdown(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", text)
    return text.replace("**", "").replace("`", "")


def configure_styles(doc: Document):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = True


def configure_page(doc: Document):
    doc.settings.odd_and_even_pages_header_footer = True
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("milXdy Roadmap Brief  |  Working Draft")
    set_run_font(run, size=9, color=MUTED, bold=True)

    even_header = section.even_page_header
    p = even_header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("milXdy Roadmap Brief  |  Working Draft")
    set_run_font(run, size=9, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.1))
    run = p.add_run("milXdy roadmap working draft - July 20, 2026")
    set_run_font(run, size=9, color=MUTED)
    p.add_run("\t")
    add_page_field(p)

    even_footer = section.even_page_footer
    p = even_footer.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.1))
    run = p.add_run("milXdy roadmap working draft - July 20, 2026")
    set_run_font(run, size=9, color=MUTED)
    p.add_run("\t")
    add_page_field(p)


def add_callout(doc, label: str, text: str, fill=PALE_GOLD):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{label}: ")
    set_run_font(run, size=10.5, color=NAVY, bold=True)
    add_inline(p, text, default_color=NAVY, default_size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("MILXDY ROADMAP BRIEF")
    set_run_font(run, size=26, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("Working roadmap for the RemiNet developer team")
    set_run_font(run, size=14, color=MUTED)

    metadata = [
        ("Status", "Content working draft - live edits expected"),
        ("Prepared", "July 20, 2026"),
        ("Audience", "RemiNet developers and adjacent Remilia ecosystem collaborators"),
        ("Scope", "0.2.x product expansion, 0.3.x Onchain Integration, and conjectural post-0.3 directions"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{label}: ")
        set_run_font(run, size=10.5, color=BLACK, bold=True)
        run = p.add_run(value)
        set_run_font(run, size=10.5, color=BLACK)

    doc.add_paragraph()
    add_callout(
        doc,
        "Editing note",
        "This draft is deliberately neutral and editable. The Remilia aesthetic pass is deferred until the narrative, scope, and diagrams are settled.",
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("DOCUMENT MAP")
    set_run_font(run, size=11, color=BLUE, bold=True)
    items = [
        "Executive direction and release map",
        "0.2.x product and RemiNet integration work",
        "0.3.x Onchain Integration sequence",
        "Conjectural research directions",
        "Editable diagrams and low-fidelity mockups",
        "Current planning positions and delivery risks",
    ]
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        add_inline(p, item)

    doc.add_page_break()


def add_roadmap_table(doc, rows):
    column_count = len(rows[0])
    table = doc.add_table(rows=0, cols=column_count)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline(p, clean_markdown(value), default_size=9.2)
    if column_count == 4:
        set_table_geometry(table, [900, 1320, 2300, 4840])
    else:
        set_table_geometry(table, [1080, 2460, 5820])
    set_repeat_table_header(table.rows[0])
    for idx, row in enumerate(table.rows):
        prevent_row_split(row)
        for cell in row.cells:
            if idx == 0:
                set_cell_shading(cell, PALE_BLUE)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = rgb(NAVY)
            elif idx % 2 == 0:
                set_cell_shading(cell, "FAFBFC")
    return table


def add_flow_box_table(doc, rows, widths, caption, fills=None):
    table = doc.add_table(rows=len(rows), cols=len(widths))
    for r_idx, row_values in enumerate(rows):
        for c_idx, value in enumerate(row_values):
            cell = table.cell(r_idx, c_idx)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_run_font(run, size=9.2, color=NAVY, bold=True)
            if fills:
                set_cell_shading(cell, fills[r_idx][c_idx])
    set_table_geometry(table, widths)
    for row in table.rows:
        prevent_row_split(row)
        for cell in row.cells:
            set_cell_margins(cell, top=70, start=120, bottom=70, end=120)
    p = doc.add_paragraph(caption, style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_mermaid_diagram(doc, code: str):
    if "Onchain Foundation" in code:
        add_flow_box_table(
            doc,
            [
                ["0.2.x\nProduct and app-platform expansion", "0.2.x\nProduct and app-platform expansion", "0.2.x\nProduct and app-platform expansion"],
                ["0.3.0\nOnchain Foundation", "0.3.0\nOnchain Foundation", "Possible 0.4.x\nExploratory directions"],
                ["0.3.1\nEthereum Media", "0.3.2\nParaclete Network", "Workspace / companion / Gotcha\nconcept research"],
                ["Collections + publishing", "Social value + BlobMail + collective metadata", "No committed order,\nversions, or scope"],
            ],
            [3120, 3120, 3120],
            "Figure 1. Product sequence. The right-hand lane is exploratory rather than a committed continuation.",
            [
                [PALE_BLUE, PALE_BLUE, PALE_BLUE],
                ["DDEAF5", "DDEAF5", PALE_GOLD],
                ["EDF4F9", "EDF4F9", "FFFDF5"],
                ["F6F9FB", "F6F9FB", "FFFDF5"],
            ],
        )
    elif "Supported RemiNet lookup" in code:
        rows = [
            ["Input", "Lookup result", "Product response"],
            ["Stable X account ID", "Eligible", "Show explicit Poke / friend action"],
            ["Stable X account ID", "Ineligible", "Hide unavailable action"],
            ["Viewer/session", "Unauthenticated", "Explain sign-in requirement"],
            ["Network", "Rate limited", "Preserve unknown; retry later"],
            ["Network/cache", "Unavailable or stale", "Preserve unknown; retry later"],
        ]
        add_flow_box_table(
            doc,
            rows,
            [2500, 2500, 4360],
            "Figure 2. RemiNet profile integration states. Failure is not treated as evidence of ineligibility.",
            [[PALE_BLUE] * 3] + [[WHITE] * 3 for _ in range(5)],
        )
    elif "User-approved links" in code:
        rows = [
            ["Identity surface", "Possible approved link", "Boundary"],
            ["X account", "RemiNet / milXdy registry", "No silent equivalence"],
            ["RemiNet", "Wallet / ENS", "Explicit proof and unlink"],
            ["Wallet", "BlobMail", "Separate key and capability model"],
            ["Observer identity", "Collective Metadata", "Provenance is not truth"],
        ]
        add_flow_box_table(
            doc,
            rows,
            [2500, 2860, 4000],
            "Figure 3. Identity systems may link, but none automatically stands in for another.",
            [[PALE_BLUE] * 3] + [[WHITE] * 3 for _ in range(4)],
        )
    elif "Canonical signed observation" in code:
        rows = [
            ["1", "Public value appears during ordinary browsing"],
            ["2", "Approved passive adapter creates canonical observation input"],
            ["3", "Observer signs a bounded content-addressed record"],
            ["4", "Post-Paraclete metadata channel propagates the record"],
            ["5", "Clients validate, deduplicate, expire, and index locally"],
            ["6", "Consumer uses the record, beginning with Account based in filtering"],
        ]
        add_flow_box_table(
            doc,
            rows,
            [900, 8460],
            "Figure 4. Collective Metadata is a bounded observation pipeline, not a generic message bus.",
            [[PALE_BLUE, WHITE] for _ in rows],
        )
    else:
        add_code_block(doc, code)


def add_code_block(doc, code: str):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(code.splitlines()):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=8.4, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_table(lines, start):
    block = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        block.append(lines[idx].strip())
        idx += 1
    rows = []
    for line_idx, line in enumerate(block):
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if line_idx == 1 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows, idx


def add_heading(doc, text, level, first_heading_state):
    p = doc.add_paragraph(style=f"Heading {level}")
    continuation_sections = {
        "Product and dependency sequence",
        "Possible 0.4.x directions: research concepts",
        "Working diagrams and low-fidelity mockups",
        "Current planning positions",
        "Major risks and gates",
        "Reference planning documents",
    }
    if level == 1 and first_heading_state[0] and clean_markdown(text) not in continuation_sections:
        p.paragraph_format.page_break_before = True
    add_inline(p, clean_markdown(text), default_color=BLUE if level < 3 else DARK_BLUE, default_size={1: 16, 2: 13, 3: 12}[level])
    first_heading_state[0] = True


def build_document():
    source = SOURCE.read_text(encoding="utf-8-sig")
    lines = source.splitlines()

    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    doc.core_properties.title = "milXdy Roadmap Brief for the RemiNet Developer Team"
    doc.core_properties.subject = "Working product roadmap and integration schedule"
    doc.core_properties.author = "milXdy"
    doc.core_properties.keywords = "milXdy, RemiNet, roadmap, integration, working draft"

    add_title_page(doc)

    idx = 0
    in_code = False
    code_lang = ""
    code_lines = []
    first_heading_state = [False]

    while idx < len(lines):
        raw = lines[idx]
        stripped = raw.strip()

        if idx == 0 and stripped.startswith("# "):
            idx += 1
            continue
        if stripped.startswith("**Status:**") or stripped.startswith("**Prepared:**") or stripped.startswith("**Audience:**") or stripped.startswith("**Repository:**"):
            idx += 1
            continue

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = stripped[3:].strip()
                code_lines = []
            else:
                code = "\n".join(code_lines)
                if code_lang == "mermaid":
                    add_mermaid_diagram(doc, code)
                else:
                    add_code_block(doc, code)
                in_code = False
                code_lang = ""
                code_lines = []
            idx += 1
            continue

        if in_code:
            code_lines.append(raw)
            idx += 1
            continue

        if not stripped:
            idx += 1
            continue

        if stripped.startswith("|"):
            rows, idx = parse_table(lines, idx)
            add_roadmap_table(doc, rows)
            continue

        heading_match = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading_match:
            hashes, title = heading_match.groups()
            add_heading(doc, title, min(len(hashes) - 1, 3), first_heading_state)
            idx += 1
            continue

        bullet_match = re.match(r"^-\s+(.+)$", stripped)
        if bullet_match:
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, bullet_match.group(1))
            idx += 1
            continue

        number_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if number_match:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, number_match.group(1))
            idx += 1
            continue

        p = doc.add_paragraph()
        add_inline(p, stripped)
        idx += 1

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    output = build_document()
    print(output)
