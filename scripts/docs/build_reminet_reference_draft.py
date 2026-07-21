from __future__ import annotations

from pathlib import Path
import json
import os
import re
import subprocess

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = Path(
    os.environ.get(
        "MILXDY_ROADMAP_DOCX_OUTPUT",
        ROOT / "ideas" / "meeting-drafts" / "milxdy-roadmap-brief-reminet-final.docx",
    )
)

PAPER = "F8F9F8"
INK = "111827"
BLUE = "2563EB"
RED = "DC2626"
SLATE = "475569"
LABEL = "F1F5F9"
BORDER = "CBD5E1"
FONT = "Roboto Mono"


def color(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_font(run, *, size=9, bold=False, italic=False, value=INK):
    run.font.name = FONT
    rfonts = run._element.get_or_add_rPr().rFonts
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color(value)


def shade(cell, fill: str):
    props = cell._tc.get_or_add_tcPr()
    node = props.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        props.append(node)
    node.set(qn("w:fill"), fill)


def borders(cell):
    props = cell._tc.get_or_add_tcPr()
    node = props.first_child_found_in("w:tcBorders")
    if node is None:
        node = OxmlElement("w:tcBorders")
        props.append(node)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = node.find(qn(f"w:{side}"))
        if edge is None:
            edge = OxmlElement(f"w:{side}")
            node.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "6")
        edge.set(qn("w:color"), BORDER)


def margins(cell, top=75, start=100, bottom=75, end=100):
    props = cell._tc.get_or_add_tcPr()
    node = props.first_child_found_in("w:tcMar")
    if node is None:
        node = OxmlElement("w:tcMar")
        props.append(node)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        child = node.find(qn(f"w:{edge}"))
        if child is None:
            child = OxmlElement(f"w:{edge}")
            node.append(child)
        child.set(qn("w:w"), str(value))
        child.set(qn("w:type"), "dxa")


def setup_table(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cell)
            borders(cell)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = "1"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_font(run, value=SLATE)


def set_background(section):
    background = OxmlElement("w:background")
    background.set(qn("w:color"), PAPER)
    section._sectPr.getparent().insert(0, background)


def para(doc, text="", *, before=0, after=5, keep=False, align=None, value=INK, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = Pt(11.87)
    p.paragraph_format.keep_with_next = keep
    if align:
        p.alignment = align
    set_font(p.add_run(text), value=value, bold=bold, italic=italic)
    return p


def heading(doc, number, text):
    return para(doc, f"{number}. {text.upper()}", before=10, after=5, keep=True, value=BLUE, bold=True)


def bullet(doc, text):
    p = para(doc, after=1)
    p.paragraph_format.left_indent = Inches(.22)
    p.paragraph_format.first_line_indent = Inches(-.16)
    set_font(p.add_run("- "), bold=True)
    set_font(p.add_run(text))
    return p


def kv_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    setup_table(table, (2805, 6555))
    for row, (label, value) in zip(table.rows, rows):
        left, right = row.cells
        shade(left, LABEL)
        left.text = right.text = ""
        set_font(left.paragraphs[0].add_run(label.upper()), bold=True, value=SLATE)
        set_font(right.paragraphs[0].add_run(value))
        for cell in (left, right):
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.line_spacing = Pt(11.87)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def release_table(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    setup_table(table, (1140, 1150, 2850, 4220))
    headers = ("RELEASE", "TARGET", "WORKING NAME", "INTENDED OUTCOME")
    for cell, text in zip(table.rows[0].cells, headers):
        shade(cell, LABEL); cell.text = ""; set_font(cell.paragraphs[0].add_run(text), bold=True, value=SLATE)
    for release, target, name, outcome in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, (release, target, name, outcome)):
            cell.text = ""; set_font(cell.paragraphs[0].add_run(text))
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.line_spacing = Pt(10.5)
    setup_table(table, (1140, 1150, 2850, 4220))


def hyperlink(paragraph, text, url):
    rel = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rel)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), FONT)
    size = OxmlElement("w:sz"); size.set(qn("w:val"), "18")
    shade_node = OxmlElement("w:color"); shade_node.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single")
    props.extend([fonts, size, shade_node, underline])
    run.append(props)
    node = OxmlElement("w:t"); node.text = text
    run.append(node); link.append(run); paragraph._p.append(link)


def roadmap_issues():
    command = [
        "gh", "issue", "list", "--repo", "bonklek/milXdy", "--state", "open", "--limit", "200",
        "--json", "number,title,milestone,url",
    ]
    raw = subprocess.run(command, check=True, capture_output=True, text=True).stdout
    items = json.loads(raw)
    selected = [
        item for item in items
        if item.get("milestone") and item["milestone"]["title"].startswith(("0.2", "0.3"))
    ]
    def milestone_key(item):
        match = re.match(r"^(\d+)\.(\d+)\.(\d+)", item["milestone"]["title"])
        if not match:
            return (999, 999, 999, item["milestone"]["title"])
        return (*map(int, match.groups()), item["milestone"]["title"])

    return sorted(selected, key=lambda item: (milestone_key(item), item["number"]))


def issue_appendix(doc):
    heading(doc, "7", "GitHub issue reference")
    para(doc, "Each entry links to its live GitHub issue. The labels are written so the appendix can be read without opening the tracker.", after=5, value=SLATE)
    current = None
    for item in roadmap_issues():
        milestone = item["milestone"]["title"]
        if milestone != current:
            current = milestone
            para(doc, current.upper(), before=7, after=2, bold=True, value=BLUE)
        p = para(doc, after=1)
        p.paragraph_format.left_indent = Inches(.22)
        p.paragraph_format.first_line_indent = Inches(-.16)
        set_font(p.add_run("- "), bold=True)
        hyperlink(p, f"#{item['number']} — {item['title']}", item["url"])


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.78)
    section.left_margin = section.right_margin = Inches(1)
    set_background(section)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(9)
    normal.font.color.rgb = color(INK)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(header.add_run("REMILIA / MILXDY    ROADMAP BRIEF"), bold=True, value=SLATE)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("20 JUL 2026 / "), value=SLATE)
    add_page_field(footer)

    para(doc, "MILXDY ROADMAP / REMINET", after=3, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    para(doc, "**********************************************", after=10, align=WD_ALIGN_PARAGRAPH.CENTER, value=SLATE)
    heading(doc, "1", "Executive summary")
    para(doc, "milXdy is becoming a user-controlled browser app platform around X, Remilia surfaces, local media and knowledge tools, and selected onchain systems.")
    bullet(doc, "0.2.x — product and platform expansion: app runtime, reviewed creation, Reader Voice, social controls, activity, identifier media, live content, and distribution groundwork.")
    bullet(doc, "0.3.x — Onchain Integration: establish wallet and transaction safety, then sequence Radio Free Ethereum, Paraclete, CULT tipping, Blobcast, collection context, publishing, and collective metadata.")
    bullet(doc, "Possible 0.4.x directions — Personal Workspace and Remilia Gotcha are research concepts, not committed releases.")
    para(doc, "Release dates are targets: a safety or dependency gate moves the release and later dates are re-baselined.", italic=True, value=SLATE)

    heading(doc, "2", "Release index")
    release_table(doc, [
        ("0.2.3", "26 Jul", "App Runtime And Distribution Prep", "Lifecycle, recovery, distribution prep, accessibility, targeted RemiNet correctness, and App SDK closeout."),
        ("0.2.4", "02 Aug", "Composer Kit", "Post upgrades, reviewed media sharing, and Miladychan posting."),
        ("0.2.5", "09 Aug", "Reader Voice", "Long-form extraction and review, TTS/audio, X DM Read Aloud, Wikipedia/Substack support, and MP3 export."),
        ("0.2.6", "16 Aug", "Social Tuning", "Maxxer customization, friend discovery, engagement presets, RemiNet profile affordances, and Miladychan navigability."),
        ("0.2.7", "23 Aug", "Activity Arcade", "Pokes, stats, daily rituals, Banners tracking, Beetle sharing, and local relationship activity."),
        ("0.2.8", "30 Aug", "Identifier Media Layer", "Radio, Miladychan board sharing, books, podcasts, movies, recipes, source archives, and evaluated RemiCast/Twitch integration."),
        ("0.2.9", "06 Sep", "Front Door & Platform Reach", "Onboarding, store preparation, Safari desktop app path, mobile research, and small-format identity assets."),
        ("0.3.0–.7", "13 Sep–01 Nov", "Onchain Integration", "Wallet foundation, media, Paraclete, CULT Cheer, mail, collections, publishing, and shared metadata."),
        ("0.4.x?", "Open", "Research directions", "Personal Workspace and local-first collector-game concepts; no fixed milestones."),
    ])
    heading(doc, "3", "0.2.x product and platform expansion")
    para(doc, "0.2.3 / APP RUNTIME AND DISTRIBUTION PREP", before=2, after=2, bold=True)
    para(doc, "Make milXdy feel steadier as people move between X and Remilia sites: fewer broken surfaces, better recovery when a page changes, smoother app access, and a Poke button that appears only when there is actually someone to Poke. Close out the App SDK so future packages have a stable platform.")
    para(doc, "0.2.4 / COMPOSER KIT", before=5, after=2, bold=True)
    para(doc, "Upgrade posting with Miladychan posting, reply media, and “milady,” “remilio,” or custom reply helpers. Improve sharing with a Tweet PNG customizer and reviewed exports; direct global-chat destinations remain research until their session boundaries are specified.")
    para(doc, "0.2.5 / READER VOICE", before=5, after=2, bold=True)
    para(doc, "Turn long articles into something you can listen to. Reader Voice pulls out readable text into an editable popout, then reads it aloud or exports MP3 through a custom TTS engine. It covers X native articles and DMs, plus Wikipedia and Substack, with local and ElevenLabs setup guidance where it can fit cleanly.")
    para(doc, "0.2.6 / SOCIAL TUNING", before=5, after=2, bold=True)
    para(doc, "Give Maxxer more personality with custom and derivative collections and emoji likes, plus friend and social-graph discovery, unfollow suggestions, usage awareness, and explicit RemiNet friend affordances. It also upgrades Miladychan navigability with thread watchlists and optional deck/live-board views: clear board/thread context, readable defaults, lightweight navigation, and graceful unavailable states without flattening the native pseudonymous culture. X follows and RemiNet friend requests remain separate things.")
    para(doc, "0.2.7 / ACTIVITY ARCADE", before=5, after=2, bold=True)
    para(doc, "Make RemiNet activity feel alive with an optional Poke feed, Poke stats, playful Poke notifications, leaderboards, daily Milady spins, a Banners daily deck tracker, local relationship activity, and Beetle Hunt catch sharing.")
    para(doc, "0.2.8 / IDENTIFIER MEDIA LAYER", before=5, after=2, bold=True)
    para(doc, "Make milXdy a place to pass around things worth returning to: personal and board-inspired radio stations, bookmarked posts, books, podcasts, movies, recipes, source archives, and Miladychan board discussion. Identifier-first design keeps discussion attached to stable references rather than hosted copyrighted media. Each item keeps its own stable reference, so a conversation can return to the same book, episode, film, recipe, or source later. RemiCast is an explicit entry point; Twitch begins as an open/handoff path while embedding remains subject to opt-in, privacy, autoplay, and performance gates.")
    para(doc, "0.2.9 / FRONT DOOR & PLATFORM REACH", before=5, after=2, bold=True)
    para(doc, "Make the extension easier to discover, understand, and install. This covers a proper first-run guide, public screenshots, Chrome Web Store preparation, a Safari desktop app path, mobile feasibility research, and the small logo/favicons that make milXdy recognizable in tight spaces.")

    heading(doc, "4", "0.3.x onchain integration")
    para(doc, "0.3.0 / THE ONCHAIN FOUNDATION", before=2, after=2, bold=True)
    para(doc, "Give people a safe, legible starting point for connecting a browser-extension wallet to milXdy: see which account is connected, what network is involved, exactly what a wallet will be asked to do, and a clear record after it happens. milXdy never asks for seed phrases or private keys.")
    para(doc, "0.3.1 / BLOBCAST", before=5, after=2, bold=True)
    para(doc, "Bring Blobcast into milXdy as a station you can discover and listen to, with the source shown clearly enough to know what you are receiving. Making a station comes later.")
    para(doc, "0.3.2 / THE PARACLETE NETWORK", before=5, after=2, bold=True)
    para(doc, "Explore Paraclete as an optional onchain network layer in the browser. The first user value is visibility: see what is happening to an account-abstraction transaction instead of having it disappear into a black box. Scope IP-privacy and other optional features jointly with Tim Clancy, without claiming privacy properties before the protocol and threat model support them.")
    para(doc, "0.3.3 / CULT CHEER", before=5, after=2, bold=True)
    para(doc, "Let people send a CULT cheer or a direct tip with the recipient, token, amount, and network plainly visible before their wallet opens. Private or shielded transfers remain separately gated research, contingent on the Paraclete protocol, wallet support, and a plain-language privacy explanation.")
    para(doc, "0.3.4 / BLOBMAIL", before=5, after=2, bold=True)
    p = para(doc, "Bring encrypted delayed mail to milXdy as a separate thing from RemiNet Chat: a message can be written, sent into a shared batch, and later carry evidence of delivery. The first version stays on testnet development networks while the protocol proves itself. ")
    hyperlink(p, "Read the BlobMail RFC", "https://github.com/bonklek/blobmail/blob/main/REQUEST_FOR_COMMENTS.md")
    para(doc, "0.3.5–0.3.7 / COLLECTIONS, PUBLISHING, COLLECTIVE METADATA", before=5, after=2, bold=True)
    para(doc, "Show optional collection context around Remilia experiences without turning milXdy into a trading screen. Later, let RFE creators publish and keep a portable station. Collective Metadata lets participating clients share a small, signed public fact—starting with X’s Account based in field—so the same public page does not need to be repeatedly scanned. It may later complement Miladychan board sharing where the public-record model is appropriate, not replace it. It also provides the right later home for opt-in milXdy registry records with explicit proof, expiry, withdrawal, and abuse rules.")

    heading(doc, "5", "Possible 0.4.x directions")
    para(doc, "This is a research horizon, not a release commitment. Version numbers, order, architecture, and inclusion remain open.", italic=True, value=SLATE)
    kv_table(doc, [
        ("Personal Workspace", "Private sticky-note/canvas research: durable local state, accessibility, persistence, recovery, portability, agent-produced podcasts, passive local X-account dossiers, and user-directed desktop-LLM-client knowledge synthesis. Exocore is possible integration context, not a commitment."),
        ("Remilia Gotcha", "A local collector game built from the REMIGACHA prototype: spot an avatar, earn an encounter, open packs, collect cards, and level duplicates. The first version is local-only—no NFT claims, marketplace, paid packs, or public leaderboard."),
    ])
    p = para(doc, "Remilia Gotcha research: ", after=8, value=SLATE)
    hyperlink(p, "#169 — GitHub backlog issue", "https://github.com/bonklek/milXdy/issues/169")
    set_font(p.add_run(" / "), value=SLATE)
    hyperlink(p, "Isabisabel Gacha reference", "https://isabisabel.com/gacha/")
    set_font(p.add_run(" / "), value=SLATE)
    hyperlink(p, "REMIGACHA prototype", "https://github.com/bonklek/REMIGACHA-prototype")

    heading(doc, "6", "Gates and open questions")
    bullet(doc, "RemiNet upload, friend, eligibility, and Poke behavior need supported contracts rather than brittle assumptions.")
    bullet(doc, "X-facing features need stable IDs, SPA-safe lifecycle handling, bounded observation, and graceful unknown states.")
    bullet(doc, "BlobMail needs canonical vectors; Paraclete needs inspectable AA state; Collective Metadata needs conflict, abuse, retention, expiry, and withdrawal controls.")
    bullet(doc, "Native messaging, Safari, and mobile each introduce specific packaging and review constraints.")
    bullet(doc, "Public data becomes more sensitive when aggregated; shared observations need privacy limits and user controls.")

    issue_appendix(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
